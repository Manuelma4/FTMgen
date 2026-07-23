# -*- coding: utf-8 -*-
"""Génération de la fiche de travaux modificative au format Word."""
from __future__ import annotations

from decimal import Decimal, InvalidOperation
from pathlib import Path
import re
from typing import Any
import unicodedata

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .relations import object_relation_key


CATEGORY_LABELS = (
    ("architect", "Adaptations demandées par le Maître d’Œuvre"),
    ("owner", "Adaptations demandées par le Maître d’Ouvrage"),
    ("program", "Changement de programme"),
    ("regulation", "Changement de réglementation"),
    ("technical", "Modification et optimisation technique"),
    ("other", "Autre cas (à préciser)"),
)

ATTACHMENT_LABELS = (
    ("plans", "Plans"),
    ("summary", "Descriptif sommaire"),
    ("estimate", "Estimation MOE"),
    ("other", "Autres"),
)

RECIPIENT_LABELS = (
    ("owner", "Maîtrise d’Ouvrage"),
    ("assistant", "Assistant Maîtrise d’Ouvrage"),
    ("company", "Entreprise"),
)


def _clean_text(value: Any, max_length: int = 4000) -> str:
    text = str(value or "").replace("\x00", "").strip()
    return text[:max_length]


def _clean_flags(value: Any, keys: tuple[str, ...]) -> dict[str, bool]:
    source = value if isinstance(value, dict) else {}
    return {key: bool(source.get(key, False)) for key in keys}


def _legacy_quantity_after(before: str, difference: str) -> str:
    """Convertit l'ancien couple quantité marché/écart vers avant/après."""
    if not difference:
        return before
    try:
        total = Decimal((before or "0").replace(" ", "").replace(",", ".")) + \
            Decimal(difference.replace(" ", "").replace(",", "."))
    except InvalidOperation:
        return ""
    rendered = format(total, "f")
    if "." in rendered:
        rendered = rendered.rstrip("0").rstrip(".")
    return rendered or "0"


def normalize_ftm_document(payload: dict[str, Any] | None) -> dict[str, Any]:
    """Nettoie les données éditées dans l'interface avant persistance/DOCX."""
    source = payload if isinstance(payload, dict) else {}
    materials: list[dict[str, str]] = []
    for raw in (source.get("materials") or [])[:250]:
        if not isinstance(raw, dict):
            continue
        quantity_before = _clean_text(raw.get("quantity_before", raw.get("market_quantity")), 40)
        if "quantity_after" in raw:
            quantity_after = _clean_text(raw.get("quantity_after"), 40)
        else:
            quantity_after = _legacy_quantity_after(
                quantity_before, _clean_text(raw.get("additional_quantity"), 40)
            )
        row = {
            "id": _clean_text(raw.get("id"), 80),
            "mapping_key": _clean_text(raw.get("mapping_key"), 500),
            "origin": "manual" if str(raw.get("origin") or "").lower() == "manual" else "pdf",
            "room": _clean_text(raw.get("room"), 240),
            "material": _clean_text(raw.get("material"), 500),
            "category": _clean_text(raw.get("category"), 240),
            "comparison_room": _clean_text(raw.get("comparison_room"), 240),
            "comparison_material": _clean_text(raw.get("comparison_material"), 500),
            "is_addition": bool(raw.get("is_addition", False)),
            "quantity_before": quantity_before,
            "quantity_after": quantity_after,
            "unit_price": _clean_text(raw.get("unit_price"), 40),
            "company_price": _clean_text(raw.get("company_price"), 40),
        }
        # Une ligne totalement vide n'a pas à apparaître dans le document final.
        if row["room"] or row["material"]:
            materials.append(row)

    return {
        "excel_scope_id": _clean_text(source.get("excel_scope_id"), 80),
        "project_name": _clean_text(source.get("project_name"), 300),
        "project_description": _clean_text(source.get("project_description"), 500),
        "issuer": _clean_text(source.get("issuer"), 120),
        "ftm_number": _clean_text(source.get("ftm_number"), 40),
        "revision": _clean_text(source.get("revision"), 20),
        "subject": _clean_text(source.get("subject"), 800),
        "pole": _clean_text(source.get("pole"), 300),
        "lot": _clean_text(source.get("lot"), 300),
        "floor": _clean_text(source.get("floor"), 120),
        "description": _clean_text(source.get("description"), 10000),
        "categories": _clean_flags(source.get("categories"), tuple(key for key, _ in CATEGORY_LABELS)),
        "category_other": _clean_text(source.get("category_other"), 500),
        "attachments": _clean_flags(source.get("attachments"), tuple(key for key, _ in ATTACHMENT_LABELS)),
        "attachment_other": _clean_text(source.get("attachment_other"), 500),
        "recipients": _clean_flags(source.get("recipients"), tuple(key for key, _ in RECIPIENT_LABELS)),
        "architect_signatory": _clean_text(source.get("architect_signatory"), 200),
        "assistant_signatory": _clean_text(source.get("assistant_signatory"), 200),
        "owner_signatory": _clean_text(source.get("owner_signatory"), 200),
        "decision": _clean_text(source.get("decision"), 30).lower(),
        "materials_version": 3,
        "materials": materials,
    }


def _norm_token(value: Any) -> str:
    text = unicodedata.normalize("NFD", str(value or ""))
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def _quantity_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return _clean_text(value, 40)


def _row_material_keys(row: dict[str, Any]) -> set[str]:
    room = _norm_token(row.get("room"))
    return {
        f"{room}|{material}"
        for material in (
            _norm_token(row.get("material")),
            _norm_token(row.get("comparison_material")),
        )
        if room or material
    }


def pdf_material_groups(
    analysis: dict[str, Any],
    include_ignored: bool = False,
) -> list[dict[str, Any]]:
    """Regroupe la traçabilité PDF sans perdre l'identité de la relation."""
    groups: dict[str, dict[str, Any]] = {}
    for trace in (analysis.get("traceabilite") or []):
        if not isinstance(trace, dict) or (trace.get("ignored") and not include_ignored):
            continue
        room = _clean_text(trace.get("room"), 240)
        material = _clean_text(
            trace.get("article") or trace.get("original_article") or trace.get("materiel_compare"), 500
        )
        comparison_material = _clean_text(
            trace.get("materiel_compare") or trace.get("article") or trace.get("original_article"), 500
        )
        category = _clean_text(trace.get("categorie"), 240)
        if not material:
            continue
        relation_key = object_relation_key(room, material)
        group_key = "|".join((relation_key, _norm_token(comparison_material), _norm_token(category)))
        if group_key in groups:
            groups[group_key]["quantity_after"] += 1
        else:
            groups[group_key] = {
                "mapping_key": relation_key,
                "origin": "pdf",
                "room": room,
                "material": material,
                "comparison_material": comparison_material,
                "category": category,
                "quantity_after": 1,
            }
    return sorted(groups.values(), key=lambda item: (
        _norm_token(item["room"]), _norm_token(item["material"]), _norm_token(item["category"]),
    ))


def all_pdf_relation_keys(analysis: dict[str, Any]) -> list[str]:
    # Les relations masquées restent dans l'univers PDF. Sinon un simple
    # enregistrement ultérieur oublierait l'exclusion et ferait réapparaître
    # silencieusement la ligne au prochain recalcul.
    return sorted({
        str(item["mapping_key"])
        for item in pdf_material_groups(analysis, include_ignored=True)
    })


def materials_detected_in_pdf(analysis: dict[str, Any], payload: dict[str, Any]) -> list[dict[str, str]]:
    """Reconstruit la table Word et conserve aussi les ajouts manuels explicites."""
    submitted = [item for item in (payload.get("materials") or []) if isinstance(item, dict)]
    submitted_by_relation = {
        str(item.get("mapping_key") or object_relation_key(item.get("room"), item.get("material"))): item
        for item in submitted
        if str(item.get("origin") or "pdf").lower() != "manual"
    }
    submitted_by_key: dict[str, dict[str, Any]] = {}
    for item in submitted:
        if str(item.get("origin") or "pdf").lower() == "manual":
            continue
        for key in _row_material_keys(item):
            submitted_by_key[key] = item
    try:
        preserve_selection = int(payload.get("materials_version") or 0) >= 2
    except (TypeError, ValueError):
        preserve_selection = False

    room_mappings = {
        _norm_token(item.get("plan")): _clean_text(item.get("room_key") or item.get("maquette"), 240)
        for item in (analysis.get("pieces_rapprochees") or [])
        if isinstance(item, dict) and _clean_text(item.get("plan"), 240)
    }

    comparison_rows = [row for row in (analysis.get("comparatif") or []) if isinstance(row, dict)]
    output: list[dict[str, str]] = []
    groups = pdf_material_groups(analysis, include_ignored=preserve_selection)
    for index, group in enumerate(groups, start=1):
        allowed_keys = _row_material_keys(group)
        previous = submitted_by_relation.get(group["mapping_key"]) \
            or next((submitted_by_key[key] for key in allowed_keys if key in submitted_by_key), None)
        if preserve_selection and previous is None:
            continue

        mapped_room = room_mappings.get(_norm_token(group["room"]), "")
        has_selected_room = previous is not None and "comparison_room" in previous
        submitted_material = _clean_text((previous or {}).get("comparison_material"), 500)
        is_addition = bool((previous or {}).get("is_addition", False))
        has_selected_material = bool(submitted_material) or is_addition
        selected_room = _clean_text((previous or {}).get("comparison_room"), 240) \
            if has_selected_room else mapped_room
        selected_material = "" if is_addition else (submitted_material or group["comparison_material"])
        candidate_pieces = {_norm_token(selected_room)} - {""}
        if not has_selected_room:
            candidate_pieces.update({
                _norm_token(group["room"]),
                _norm_token(f"{group['room']} [nouvelle pièce]"),
            } - {""})
        candidate_materials = {_norm_token(selected_material)} - {""}
        if not has_selected_material:
            candidate_materials.add(_norm_token(group["material"]))
            candidate_materials.discard("")
        matching = [
            row for row in comparison_rows
            if (_clean_text(row.get("room_id"), 240) == selected_room
                or _norm_token(row.get("piece")) in candidate_pieces)
            and _norm_token(row.get("materiel")) in candidate_materials
        ]
        comparison = next(
            (row for row in matching if _norm_token(row.get("categorie")) == _norm_token(group["category"])),
            matching[0] if matching else None,
        )
        output.append({
            "id": _clean_text((previous or {}).get("id"), 80) or f"pdf-{index}",
            "mapping_key": group["mapping_key"],
            "origin": "pdf",
            "room": group["room"],
            "material": group["material"],
            "category": group["category"],
            "comparison_room": selected_room,
            "comparison_material": selected_material,
            "is_addition": is_addition,
            "quantity_before": _quantity_text(comparison.get("quantite_avant")) if comparison else (
                "0" if has_selected_room or has_selected_material else ""
            ),
            "quantity_after": str(group["quantity_after"]),
            "unit_price": _clean_text((previous or {}).get("unit_price"), 40),
            "company_price": _clean_text((previous or {}).get("company_price"), 40),
        })

    for index, previous in enumerate(submitted, start=1):
        if str(previous.get("origin") or "pdf").lower() != "manual":
            continue
        room = _clean_text(previous.get("room"), 240)
        material = _clean_text(previous.get("material"), 500)
        if not room and not material:
            continue
        selected_room = _clean_text(previous.get("comparison_room"), 240)
        selected_material = _clean_text(previous.get("comparison_material"), 500)
        is_addition = bool(previous.get("is_addition", not selected_material))
        if is_addition:
            selected_material = ""
        matching = next((
            row for row in comparison_rows
            if (_clean_text(row.get("room_id"), 240) == selected_room
                or _norm_token(row.get("piece")) == _norm_token(selected_room))
            and _norm_token(row.get("materiel")) == _norm_token(selected_material)
        ), None)
        output.append({
            "id": _clean_text(previous.get("id"), 80) or f"manual-{index}",
            "mapping_key": _clean_text(previous.get("mapping_key"), 500) or f"manual-{index}",
            "origin": "manual",
            "room": room,
            "material": material,
            "category": _clean_text(previous.get("category"), 240),
            "comparison_room": selected_room,
            "comparison_material": selected_material,
            "is_addition": is_addition,
            "quantity_before": _quantity_text(matching.get("quantite_avant")) if matching else "0",
            "quantity_after": _quantity_text(previous.get("quantity_after")) or "1",
            "unit_price": _clean_text(previous.get("unit_price"), 40),
            "company_price": _clean_text(previous.get("company_price"), 40),
        })
    return output


def _set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_row_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _set_cell_width(cell, width) -> None:
    cell.width = width
    properties = cell._tc.get_or_add_tcPr()
    tc_width = properties.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        properties.append(tc_width)
    tc_width.set(qn("w:w"), str(int(width.twips)))
    tc_width.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = "000000", size: str = "8") -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def _set_run(run, *, bold: bool = False, size: float = 8, color: str = "000000") -> None:
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _write_cell(cell, text: str, *, bold: bool = False, size: float = 8,
                align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    _set_run(run, bold=bold, size=size)


def _add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend((begin, code, separate, value, end))


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(1.05)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.4)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(8)
    style.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    _set_run(run, size=8, color="777777")
    _add_field(footer, "PAGE")
    run = footer.add_run(" / ")
    _set_run(run, size=8, color="777777")
    _add_field(footer, "NUMPAGES")


def _add_heading_box(document: Document, data: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, color="137333", size="12")
    cell = table.cell(0, 0)
    _set_cell_width(cell, Cm(18.85))
    cell.text = ""
    _set_cell_margins(cell, top=120, bottom=120)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, (text, bold, size) in enumerate((
        (data["project_name"], True, 10),
        (data["project_description"], True, 9),
        ("FICHE DE TRAVAUX MODIFICATIVE", True, 11),
    )):
        if not text:
            continue
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(text)
        _set_run(run, bold=bold, size=size)


def _add_document_identity(document: Document, data: dict[str, Any]) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    widths = (Cm(4.7), Cm(4.2), Cm(2.0), Cm(2.1), Cm(3.8))
    values = (
        f"Émetteur :\n{data['issuer']}" if data["issuer"] else "Émetteur :",
        "",
        "FTM n°",
        data["ftm_number"],
        f"N° indice : {data['revision']}" if data["revision"] else "N° indice :",
    )
    for cell, width, value in zip(table.rows[0].cells, widths, values):
        _set_cell_width(cell, width)
        _write_cell(cell, value, bold=value in ("FTM n°",), align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_notice(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(
        "Le présent document fait état de la décision du maître d’ouvrage d’une intention de modifier "
        "les prestations du MARCHÉ ; en conséquence, la maîtrise d’œuvre donne l’ordre à l’entreprise "
        "de fournir sous huitaine un devis relatif à la présente Fiche de Travaux Modificative."
    )
    _set_run(run, size=7, color="555555")
    run.italic = True


def _add_general_information(document: Document, data: dict[str, Any]) -> None:
    labels = (
        ("Objet de cette Fiche de Travaux Modificative", data["subject"]),
        ("Pôle", data["pole"]),
        ("LOT", data["lot"]),
        ("Étage", data["floor"]),
    )
    table = document.add_table(rows=len(labels), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    for row, (label, value) in zip(table.rows, labels):
        _set_row_cant_split(row)
        cell = row.cells[0]
        _set_cell_width(cell, Cm(18.85))
        cell.text = ""
        _set_cell_margins(cell, top=75, bottom=75)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{label} :")
        _set_run(run, bold=True, size=8)
        if value:
            run = paragraph.add_run(f" {value}")
            _set_run(run, size=8)


def _add_materials(document: Document, materials: list[dict[str, str]]) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    title = document.add_table(rows=1, cols=1)
    title.alignment = WD_TABLE_ALIGNMENT.CENTER
    title.autofit = False
    _set_table_borders(title)
    _set_cell_width(title.cell(0, 0), Cm(18.85))
    _write_cell(
        title.cell(0, 0), "OBJETS IDENTIFIÉS SUR LE PLAN PDF",
        bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    widths = (Cm(3.0), Cm(6.0), Cm(2.15), Cm(2.15), Cm(2.35), Cm(3.2))
    headers = (
        "Nom de la pièce", "Matériel", "Quantité marché", "Quantité après FTM",
        "Prix\nunitaire", "Prix\nentreprise",
    )
    header = table.rows[0]
    _set_repeat_table_header(header)
    for cell, width, label in zip(header.cells, widths, headers):
        _set_cell_width(cell, width)
        _set_cell_shading(cell, "F2F2F2")
        _write_cell(cell, label, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)

    display_rows = materials or [{key: "" for key in (
        "room", "material", "quantity_before", "quantity_after", "unit_price", "company_price"
    )}]
    for material in display_rows:
        row = table.add_row()
        _set_row_cant_split(row)
        values = (
            material.get("room", ""), material.get("material", ""),
            material.get("quantity_before", ""), material.get("quantity_after", ""),
            material.get("unit_price", ""), material.get("company_price", ""),
        )
        for index, (cell, width, value) in enumerate(zip(row.cells, widths, values)):
            _set_cell_width(cell, width)
            _write_cell(
                cell, value, size=7,
                align=WD_ALIGN_PARAGRAPH.LEFT if index < 2 else WD_ALIGN_PARAGRAPH.CENTER,
            )
        row.height = Cm(0.75)


def _add_section_title(document: Document, title: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    _set_cell_width(table.cell(0, 0), Cm(18.85))
    _set_cell_shading(table.cell(0, 0), "F2F2F2")
    _write_cell(table.cell(0, 0), title.upper(), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_modifications_page(document: Document, data: dict[str, Any]) -> None:
    document.add_page_break()
    _add_section_title(document, "Modifications")
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_width(cell, Cm(18.85))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.text = ""
    _set_cell_margins(cell, top=100, bottom=100)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Descriptif :")
    _set_run(run, bold=True, size=8)
    if data["description"]:
        run = paragraph.add_run(f"\n\n{data['description']}")
        _set_run(run, size=8)
    row = table.rows[0]
    row.height = Cm(23.5)


def _add_checkbox_lines(document: Document, title: str, items, flags: dict[str, bool], other_text: str = "") -> None:
    _add_section_title(document, title)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_width(cell, Cm(18.85))
    cell.text = ""
    _set_cell_margins(cell, top=100, bottom=100)
    for index, (key, label) in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        suffix = f" : {other_text}" if key == "other" and other_text else ""
        run = paragraph.add_run(f"{'☒' if flags.get(key) else '☐'}  {label}{suffix}")
        _set_run(run, size=8)


def _add_administration_page(document: Document, data: dict[str, Any]) -> None:
    document.add_page_break()
    _add_checkbox_lines(document, "Catégorie de la demande", CATEGORY_LABELS, data["categories"], data["category_other"])
    document.add_paragraph().paragraph_format.space_after = Pt(6)
    _add_checkbox_lines(document, "Documents MOE joints", ATTACHMENT_LABELS, data["attachments"], data["attachment_other"])
    document.add_paragraph().paragraph_format.space_after = Pt(8)
    _add_checkbox_lines(document, "Diffusion du présent document", RECIPIENT_LABELS, data["recipients"])
    document.add_paragraph().paragraph_format.space_after = Pt(8)

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table)
    widths = (Cm(6.28), Cm(6.28), Cm(6.29))
    headings = (
        "Établie par le Maître d’Œuvre",
        "Visée par l’Assistant Maître d’Ouvrage",
        "Acceptée par le Maître de l’Ouvrage",
    )
    names = (data["architect_signatory"], data["assistant_signatory"], data["owner_signatory"])
    for index, (cell, width, heading, name) in enumerate(zip(table.rows[0].cells, widths, headings, names)):
        _set_cell_width(cell, width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.text = ""
        _set_cell_margins(cell, top=90, bottom=90)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(heading)
        _set_run(run, bold=True, size=7.5)
        if name:
            run = paragraph.add_run(f" : {name}")
            _set_run(run, bold=True, size=7.5)
        run = paragraph.add_run("\nNom, date, visa, signature")
        _set_run(run, size=7, color="555555")
        if index == 2:
            accepted = data["decision"] == "accepted"
            refused = data["decision"] == "refused"
            run = paragraph.add_run(f"\n\n{'☒' if accepted else '☐'} Acceptée\n{'☒' if refused else '☐'} Refusée")
            _set_run(run, size=8)
    table.rows[0].height = Cm(7.0)


def write_ftm_document(path: str | Path, payload: dict[str, Any]) -> dict[str, Any]:
    """Écrit le DOCX et renvoie les données normalisées effectivement utilisées."""
    data = normalize_ftm_document(payload)
    document = Document()
    _configure_document(document)
    _add_heading_box(document, data)
    _add_document_identity(document, data)
    _add_notice(document)
    _add_general_information(document, data)
    _add_materials(document, data["materials"])
    _add_modifications_page(document, data)
    _add_administration_page(document, data)

    properties = document.core_properties
    properties.title = "Fiche de Travaux Modificative"
    properties.subject = data["subject"]
    properties.author = data["issuer"] or "FTMgen"
    properties.comments = "Document généré par FTMgen à partir des informations validées dans l'interface."

    output = Path(path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return data
